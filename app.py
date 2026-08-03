import os
import io
import glob
import tempfile
import logging
import traceback
import shutil
import json
import sys
import types
sys.modules['torchvision._meta_registrations'] = types.ModuleType('torchvision._meta_registrations')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from src.core.corrector import correct_vietnamese_text
try:
    import fitz
except ImportError:
    fitz = None

import asyncio

try:
    import gradio as gr
    HAS_GRADIO = True
except ImportError:
    HAS_GRADIO = False

from typing import Optional

try:
    from fastapi import FastAPI, File, UploadFile, Query, HTTPException
    from fastapi.responses import JSONResponse, HTMLResponse, StreamingResponse
    from fastapi.middleware.cors import CORSMiddleware
    HAS_FASTAPI = True
except ImportError:
    HAS_FASTAPI = False
    def DummyFunc(*args, **kwargs):
        return None
    File = UploadFile = Query = HTTPException = DummyFunc
    JSONResponse = HTMLResponse = StreamingResponse = CORSMiddleware = None

try:
    import torch
except ImportError:
    torch = None

from concurrent.futures import ThreadPoolExecutor

if torch and torch.cuda.is_available():
    try:
        torch.set_float32_matmul_precision('high')
        torch.backends.cuda.matmul.allow_tf32 = True
        torch.backends.cudnn.allow_tf32 = True
        torch.backends.cudnn.benchmark = True
    except Exception:
        pass

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger("unlimited_ocr_app")

if HAS_FASTAPI:
    app = FastAPI(
        title="Unlimited-OCR API Server with Unsloth / PyTorch GPU",
        description="High-performance OCR & Document Parsing API hosting Baidu Unlimited-OCR",
        version="1.0.0"
    )

    # Enable CORS for frontend clients
    app.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )
else:
    class DummyApp:
        def post(self, *args, **kwargs):
            return lambda fn: fn
        def get(self, *args, **kwargs):
            return lambda fn: fn
        def on_event(self, *args, **kwargs):
            return lambda fn: fn
        def add_middleware(self, *args, **kwargs):
            pass
    app = DummyApp()

MODEL_NAME = os.getenv("MODEL_NAME", "baidu/Unlimited-OCR")
tokenizer = None
model = None
load_status = {"status": "not_loaded", "error": None, "backend": None}

def get_model():
    global tokenizer, model, load_status
    if model is not None and tokenizer is not None:
        return model, tokenizer

    logger.info(f"Initializing Unlimited-OCR model: {MODEL_NAME}...")
    load_status["status"] = "loading"
    
    try:
        from transformers import AutoModel, AutoTokenizer

        tokenizer = AutoTokenizer.from_pretrained(MODEL_NAME, trust_remote_code=True)
        logger.info("Tokenizer loaded successfully.")
        
        # Try loading via Unsloth if explicitly enabled
        if os.getenv("ENABLE_UNSLOTH", "0") == "1":
            try:
                from unsloth import FastVisionModel
                logger.info("Attempting loading via Unsloth FastVisionModel...")
                model, tokenizer = FastVisionModel.from_pretrained(
                    MODEL_NAME,
                    load_in_4bit=True,
                    trust_remote_code=True,
                )
                FastVisionModel.for_inference(model)
                load_status = {"status": "ready", "backend": "Unsloth 4-bit", "error": None}
                logger.info("Successfully initialized model using Unsloth 4-bit acceleration!")
                return model, tokenizer
            except Exception as unsloth_err:
                logger.warning(f"Unsloth initialization skipped/failed ({unsloth_err}). Falling back to standard PyTorch Transformers...")

        # Fallback to PyTorch AutoModel
        model = AutoModel.from_pretrained(
            MODEL_NAME,
            trust_remote_code=True,
            use_safetensors=True,
            torch_dtype=torch.bfloat16 if (torch.cuda.is_available() and torch.cuda.is_bf16_supported()) else torch.float16,
        )
        if torch.cuda.is_available():
            model = model.eval().cuda()
            torch.set_float32_matmul_precision('high')
            torch.backends.cuda.matmul.allow_tf32 = True
            torch.backends.cudnn.allow_tf32 = True
            torch.backends.cudnn.benchmark = True
            logger.info(f"Model loaded on GPU: {torch.cuda.get_device_name(0)}")
            load_status = {"status": "ready", "backend": "PyTorch GPU", "error": None}
        else:
            model = model.eval()
            logger.warning("CUDA unavailable, model running in CPU mode.")
            load_status = {"status": "ready", "backend": "PyTorch CPU", "error": None}

        return model, tokenizer

    except Exception as e:
        err_msg = f"{str(e)}\n{traceback.format_exc()}"
        logger.error(f"Failed to load model: {err_msg}")
        load_status = {"status": "error", "backend": None, "error": str(e)}
        raise RuntimeError(f"Model loading failed: {str(e)}")

@app.on_event("startup")
def startup_event():
    logger.info("FastAPI application started. Pre-loading model in background...")
    try:
        get_model()
    except Exception as e:
        logger.error(f"Startup model pre-load encountered error: {e}")


@app.get("/api/info")
def read_info():
    return {
        "service": "Unlimited-OCR API Server",
        "model": MODEL_NAME,
        "model_load_status": load_status,
        "gpu_available": torch.cuda.is_available(),
        "gpu_name": torch.cuda.get_device_name(0) if torch.cuda.is_available() else "None"
    }

@app.get("/health")
def health_check():
    return {"status": "healthy", "load_status": load_status}

def render_single_page(pdf_path: str, page_num: int, dpi: int = 220) -> tuple:
    doc = fitz.open(pdf_path)
    page = doc[page_num]
    pix = page.get_pixmap(dpi=dpi)
    page_tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f"_p{page_num+1}.png")
    pix.save(page_tmp.name)
    page_path = page_tmp.name
    doc.close()
    return page_num, page_path

def run_ocr_on_single_image(tmp_img_path, mode, max_length, active_model, active_tokenizer):
    output_dir = tempfile.mkdtemp(prefix="ocr_output_")
    import gc
    gc.collect()
    if torch.cuda.is_available():
        torch.cuda.empty_cache()
    try:
        crop_mode = True if mode == "gundam" else False
        image_size = 640 if mode == "gundam" else 1024
        
        with torch.inference_mode():
            res_text = active_model.infer(
                active_tokenizer,
                prompt='<image>document parsing.',
                image_file=tmp_img_path,
                output_path=output_dir,
                base_size=1024,
                image_size=image_size,
                crop_mode=crop_mode,
                max_length=max_length,
                no_repeat_ngram_size=35,
                ngram_window=128 if mode == "gundam" else 1024,
                save_results=True,
                eval_mode=True,
            )

        ocr_text = ""
        if res_text and isinstance(res_text, str) and res_text.strip():
            ocr_text = res_text.strip()
        else:
            result_md_path = os.path.join(output_dir, "result.md")
            if os.path.exists(result_md_path):
                with open(result_md_path, "r", encoding="utf-8") as f:
                    ocr_text = f.read()
            else:
                all_files = glob.glob(os.path.join(output_dir, "**", "*.*"), recursive=True)
                text_files = [f for f in all_files if f.endswith(('.md', '.txt'))]
                if text_files:
                    with open(text_files[0], "r", encoding="utf-8") as f:
                        ocr_text = f.read()
        return ocr_text.strip()
    finally:
        if os.path.exists(output_dir):
            shutil.rmtree(output_dir, ignore_errors=True)
        gc.collect()
        if torch.cuda.is_available():
            torch.cuda.empty_cache()

def format_table_line(text: str) -> str:
    import re
    if not text:
        return ""
    tokens = [t.strip() for t in re.split(r'\s{2,}|\t+', text) if t.strip()]
    if len(tokens) >= 3 and any(re.search(r'\d', t) for t in tokens):
        return "| " + " | ".join(tokens) + " |"
    return text

def clean_ocr_to_markdown(raw_text: str, page_image_path: str = None) -> str:
    import re
    if not raw_text:
        return ""
    
    # Apply administrative text corrections
    corrected_raw = correct_vietnamese_text(raw_text)
    
    lines = corrected_raw.splitlines()
    clean_lines = []
    
    for line in lines:
        match = re.match(r'^\s*<\|det\|>\s*([a-zA-Z_]+)\s*\[.*?\]\s*<\|/det\|>\s*(.*)$', line)
        if match:
            category, content = match.group(1), match.group(2).strip()
            if not content:
                continue
            
            if category == 'title':
                clean_lines.append(f"# {content}")
            elif category == 'header':
                clean_lines.append(f"### {content}")
            elif category in ['section_header', 'sub_title']:
                clean_lines.append(f"#### {content}")
            elif category == 'page_number':
                clean_lines.append(f"\n*— Trang {content} —*\n")
            else:
                formatted = format_table_line(content)
                clean_lines.append(formatted)
        else:
            cleaned = re.sub(r'<\|det\|>.*?<\|/det\|>', '', line).strip()
            if cleaned:
                formatted = format_table_line(cleaned)
                clean_lines.append(formatted)
                
    return "\n\n".join(clean_lines)

@app.post("/v1/ocr/stream")
async def process_ocr_stream(
    file: UploadFile = File(...),
    mode: str = Query("gundam"),
    max_length: int = Query(32768)
):
    active_model, active_tokenizer = get_model()
    content = await file.read()
    filename = file.filename or "uploaded_file"
    is_pdf = content.startswith(b"%PDF-") or filename.lower().endswith(".pdf")

    ext = ".pdf" if is_pdf else ".png"
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        async def event_generator():
            start_time = asyncio.get_event_loop().time()
            try:
                if is_pdf:
                    doc = fitz.open(tmp_path)
                    total_pages = len(doc)
                    doc.close()

                    # Pre-render all PDF pages concurrently across CPU threads
                    workers = min(8, max(1, total_pages))
                    loop = asyncio.get_event_loop()
                    with ThreadPoolExecutor(max_workers=workers) as executor:
                        tasks = [
                            loop.run_in_executor(executor, render_single_page, tmp_path, p_idx, 220)
                            for p_idx in range(total_pages)
                        ]
                        page_images = await asyncio.gather(*tasks)

                    page_images.sort(key=lambda x: x[0])

                    for i, page_path in page_images:
                        try:
                            raw_text = await asyncio.to_thread(run_ocr_on_single_image, page_path, mode, max_length, active_model, active_tokenizer)
                            clean_md = clean_ocr_to_markdown(raw_text, page_path)
                            elapsed = round(asyncio.get_event_loop().time() - start_time, 2)
                            data = {
                                "page_index": i + 1,
                                "total_pages": total_pages,
                                "raw_text": raw_text,
                                "clean_markdown": clean_md,
                                "elapsed_seconds": elapsed,
                                "is_complete": False
                            }
                            yield f"event: page_data\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"
                        finally:
                            if os.path.exists(page_path):
                                os.remove(page_path)
                else:
                    raw_text = await asyncio.to_thread(run_ocr_on_single_image, tmp_path, mode, max_length, active_model, active_tokenizer)
                    clean_md = clean_ocr_to_markdown(raw_text, tmp_path)
                    elapsed = round(asyncio.get_event_loop().time() - start_time, 2)
                    data = {
                        "page_index": 1,
                        "total_pages": 1,
                        "raw_text": raw_text,
                        "clean_markdown": clean_md,
                        "elapsed_seconds": elapsed,
                        "is_complete": True
                    }
                    yield f"event: page_data\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"

                total_elapsed = round(asyncio.get_event_loop().time() - start_time, 2)
                yield f"event: complete\ndata: {json.dumps({'status': 'finished', 'total_seconds': total_elapsed})}\n\n"
            finally:
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)

        return StreamingResponse(event_generator(), media_type="text/event-stream")
    except Exception:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)
        raise

@app.post("/v1/ocr")
async def process_ocr(
    file: UploadFile = File(...),
    mode: str = Query("gundam", description="Parsing mode: 'gundam' (640px crop mode) or 'base' (1024px full mode)"),
    max_length: int = Query(32768, description="Maximum output token length")
):
    filename = file.filename or "uploaded_file"
    filename_lower = filename.lower()

    try:
        active_model, active_tokenizer = get_model()
    except Exception as err:
        raise HTTPException(status_code=500, detail=f"Model failed to load: {str(err)}")

    content = await file.read()
    
    is_pdf = content.startswith(b"%PDF-") or filename_lower.endswith(".pdf") or (file.content_type in ["application/pdf", "application/x-pdf"])

    if not is_pdf and file.content_type and not file.content_type.startswith("image/") and file.content_type != "application/octet-stream":
        raise HTTPException(status_code=400, detail="Provided file must be an image or a PDF.")

    ext = ".pdf" if is_pdf else (os.path.splitext(filename)[1] or ".jpg")
    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp_file:
        tmp_file.write(content)
        tmp_file_path = tmp_file.name

    try:
        if is_pdf:
            logger.info(f"Processing PDF document for OCR: {filename} (mode: {mode})...")
            doc = fitz.open(tmp_file_path)
            total_pages = len(doc)
            logger.info(f"PDF opened successfully. Total pages: {total_pages}")
            
            page_results = []
            for i in range(total_pages):
                logger.info(f"Rendering PDF page {i+1}/{total_pages} to image...")
                page = doc[i]
                pix = page.get_pixmap(dpi=220)
                
                with tempfile.NamedTemporaryFile(delete=False, suffix=f"_page_{i+1}.png") as tmp_page_img:
                    pix.save(tmp_page_img.name)
                    tmp_page_path = tmp_page_img.name

                try:
                    page_text = await asyncio.to_thread(run_ocr_on_single_image, tmp_page_path, mode, max_length, active_model, active_tokenizer)
                    if total_pages > 1:
                        page_results.append(f"--- Page {i+1} ---\n{page_text}")
                    else:
                        page_results.append(page_text)
                finally:
                    if os.path.exists(tmp_page_path):
                        os.remove(tmp_page_path)

            final_ocr_text = "\n\n".join(page_results).strip()
            if not final_ocr_text:
                final_ocr_text = "No text extracted from PDF."
        else:
            logger.info(f"Running OCR inference on image: {filename} (mode: {mode})...")
            final_ocr_text = await asyncio.to_thread(run_ocr_on_single_image, tmp_file_path, mode, max_length, active_model, active_tokenizer)
            if not final_ocr_text:
                final_ocr_text = "No text extracted or output file not generated."

        logger.info(f"Extracted OCR text length: {len(final_ocr_text)}")
        clean_markdown_text = clean_ocr_to_markdown(final_ocr_text)

        return {
            "status": "success",
            "filename": filename,
            "is_pdf": is_pdf,
            "mode": mode,
            "parsed_text": final_ocr_text,
            "clean_markdown": clean_markdown_text
        }

    except Exception as e:
        logger.error(f"Error during OCR processing: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"OCR execution failed: {str(e)}")
    finally:
        if os.path.exists(tmp_file_path):
            os.remove(tmp_file_path)

try:
    from pydantic import BaseModel
    class ExportDocxRequest(BaseModel):
        markdown_text: str
        filename: Optional[str] = "ocr_result.docx"
except ImportError:
    class ExportDocxRequest:
        pass

if HAS_FASTAPI:
    @app.post("/v1/export/docx")
    async def export_docx_endpoint(req: ExportDocxRequest):
        if not req.markdown_text:
            raise HTTPException(status_code=400, detail="markdown_text is empty")
        try:
            import docx
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
        except ImportError:
            raise HTTPException(status_code=500, detail="python-docx library not installed on server.")

        doc = docx.Document()
        for s in doc.sections:
            s.top_margin = Inches(0.8)
            s.bottom_margin = Inches(0.8)
            s.left_margin = Inches(0.8)
            s.right_margin = Inches(0.8)

        lines = req.markdown_text.splitlines()
        in_table = False
        table_rows = []

        def flush_table():
            nonlocal in_table, table_rows
            if not table_rows:
                return
            max_cols = max(len(r) for r in table_rows)
            if max_cols > 0:
                tbl = doc.add_table(rows=len(table_rows), cols=max_cols)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl.style = 'Table Grid'
                for r_idx, row in enumerate(table_rows):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < max_cols:
                            cell = tbl.cell(r_idx, c_idx)
                            cell.text = cell_text.strip()
                            if r_idx == 0:
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:val'), 'clear')
                                shd.set(qn('w:color'), 'auto')
                                shd.set(qn('w:fill'), '1E40AF')
                                cell._tc.get_or_add_tcPr().append(shd)
                                for p in cell.paragraphs:
                                    for run in p.runs:
                                        run.font.color.rgb = RGBColor(255, 255, 255)
                                        run.font.bold = True
            table_rows = []
            in_table = False
            doc.add_paragraph()

        for line in lines:
            sline = line.strip()
            if sline.startswith('|') and sline.endswith('|'):
                parts = [p.strip() for p in sline.split('|')[1:-1]]
                if all(set(p) <= {'-', ':', ' '} for p in parts if p):
                    continue
                in_table = True
                table_rows.append(parts)
            else:
                if in_table:
                    flush_table()
                if not sline:
                    continue
                if sline.startswith('# '):
                    doc.add_heading(sline[2:], level=1)
                elif sline.startswith('## '):
                    doc.add_heading(sline[3:], level=2)
                elif sline.startswith('### '):
                    doc.add_heading(sline[4:], level=3)
                elif sline.startswith('#### '):
                    doc.add_heading(sline[5:], level=4)
                elif sline.startswith('- ') or sline.startswith('* '):
                    doc.add_paragraph(sline[2:], style='List Bullet')
                else:
                    doc.add_paragraph(sline)

        if in_table:
            flush_table()

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        
        out_filename = req.filename if req.filename and req.filename.endswith('.docx') else f"{req.filename or 'ocr_result'}.docx"
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={out_filename}"}
        )

    @app.post("/v1/import/file")
    async def import_file_endpoint(file: UploadFile = File(...)):
        filename = file.filename or "imported_file"
        filename_lower = filename.lower()
        content_bytes = await file.read()
        
        markdown_result = ""
        
        if filename_lower.endswith((".txt", ".md")):
            markdown_result = content_bytes.decode("utf-8", errors="ignore")
        elif filename_lower.endswith(".docx"):
            try:
                import docx
                doc = docx.Document(io.BytesIO(content_bytes))
                paragraphs = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        if p.style.name.startswith("Heading 1"):
                            paragraphs.append(f"# {p.text}")
                        elif p.style.name.startswith("Heading 2"):
                            paragraphs.append(f"## {p.text}")
                        elif p.style.name.startswith("Heading 3"):
                            paragraphs.append(f"### {p.text}")
                        else:
                            paragraphs.append(p.text)
                for table in doc.tables:
                    table_md = []
                    for r_idx, row in enumerate(table.rows):
                        row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                        table_md.append("| " + " | ".join(row_cells) + " |")
                        if r_idx == 0:
                            table_md.append("| " + " | ".join(["---"] * len(row_cells)) + " |")
                    paragraphs.append("\n" + "\n".join(table_md) + "\n")
                markdown_result = "\n\n".join(paragraphs)
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Không thể đọc file Word (.docx): {str(e)}")
        elif filename_lower.endswith(".pdf"):
            try:
                import fitz
                doc = fitz.open(stream=content_bytes, filetype="pdf")
                page_texts = []
                for i, page in enumerate(doc):
                    text = page.get_text("text")
                    page_texts.append(f"### --- Trang {i+1} ---\n\n{text}")
                markdown_result = "\n\n".join(page_texts)
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Không thể đọc file PDF: {str(e)}")
        else:
            raise HTTPException(status_code=400, detail="Định dạng không hỗ trợ. Vui lòng nộp file .docx, .pdf, .txt hoặc .md")

        clean_md = clean_ocr_to_markdown(markdown_result)
        return {
            "status": "success",
            "filename": filename,
            "markdown_text": clean_md,
            "raw_text": markdown_result
        }

if HAS_GRADIO:
    def gradio_ocr_process(file_obj, mode_choice):
        if file_obj is None:
            return "❌ Please upload a document image or PDF file.", ""
        
        file_path = file_obj.name
        file_name = os.path.basename(file_path)
        filename_lower = file_name.lower()
        is_pdf = filename_lower.endswith(".pdf")

        try:
            active_model, active_tokenizer = get_model()
        except Exception as err:
            return f"❌ Model loading failed: {str(err)}", ""

        if is_pdf:
            doc = fitz.open(file_path)
            total_pages = len(doc)
            page_results = []
            for i in range(total_pages):
                page = doc[i]
                pix = page.get_pixmap(dpi=150)
                with tempfile.NamedTemporaryFile(delete=False, suffix=f"_page_{i+1}.png") as tmp_page_img:
                    pix.save(tmp_page_img.name)
                    tmp_page_path = tmp_page_img.name
                try:
                    page_text = run_ocr_on_single_image(tmp_page_path, mode_choice, 32768, active_model, active_tokenizer)
                    if total_pages > 1:
                        page_results.append(f"### --- Page {i+1} ---\n\n{page_text}")
                    else:
                        page_results.append(page_text)
                finally:
                    if os.path.exists(tmp_page_path):
                        os.remove(tmp_page_path)
            final_text = "\n\n".join(page_results).strip()
        else:
            final_text = run_ocr_on_single_image(file_path, mode_choice, 32768, active_model, active_tokenizer)

        return final_text, final_text

    with gr.Blocks(title="Baidu Unlimited-OCR Web Demo", theme=gr.themes.Soft()) as demo:
        gr.Markdown("# 🚀 Baidu Unlimited-OCR Web Interface")
        gr.Markdown("Upload any Document Image or PDF file for **high-precision OCR and document layout parsing** powered by GPU acceleration.")
        
        with gr.Row():
            with gr.Column(scale=1):
                file_input = gr.File(label="📄 Upload Document Image or PDF File", file_types=["image", ".pdf"])
                mode_input = gr.Radio(
                    choices=[("Gundam Mode (640px Crop - Detailed)", "gundam"), ("Base Mode (1024px Full Page)", "base")],
                    value="gundam",
                    label="⚙️ Parsing Mode"
                )
                submit_btn = gr.Button("⚡ Run OCR & Parse Document", variant="primary")
                
            with gr.Column(scale=2):
                with gr.Tabs():
                    with gr.TabItem("📝 Extracted Text / Markdown"):
                        text_output = gr.Textbox(label="Raw Markdown Text", lines=20, show_copy_button=True)
                    with gr.TabItem("👁️ Formatted Preview"):
                        markdown_preview = gr.Markdown(label="Rendered Document Preview")

        submit_btn.click(
            fn=gradio_ocr_process,
            inputs=[file_input, mode_input],
            outputs=[text_output, markdown_preview]
        )

    app = gr.mount_gradio_app(app, demo, path="/gradio")
else:
    @app.get("/", response_class=HTMLResponse)
    def web_ui_index():
        html_content = r"""<!DOCTYPE html>
<html lang="vi">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Unlimited-OCR Studio | Baidu AI GPU</title>
    <script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/html2pdf.js/0.10.1/html2pdf.bundle.min.js"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap" rel="stylesheet">
    <style>
        :root {
            --bg-page: #f8fafc;
            --bg-card: #ffffff;
            --text-main: #0f172a;
            --text-sub: #64748b;
            --primary: #10b981;
            --primary-hover: #059669;
            --primary-dark: #047857;
            --primary-light: #ecfdf5;
            --border: #e2e8f0;
            --border-hover: #cbd5e1;
            --shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.04);
            --radius: 10px;
        }

        * {
            box-sizing: border-box;
            margin: 0;
            padding: 0;
            font-family: -apple-system, BlinkMacSystemFont, 'Inter', 'Segoe UI', Roboto, sans-serif;
        }

        body {
            background-color: var(--bg-page);
            color: var(--text-main);
            min-height: 100vh;
            display: flex;
            flex-direction: column;
        }

        main {
            max-width: 1600px;
            width: 100%;
            margin: 12px auto;
            padding: 0 16px;
            flex: 1;
        }

        /* TOP COMPACT CONTROL BAR (NOTION STYLE) */
        .top-control-bar {
            background: #ffffff;
            border: 1px solid var(--border);
            border-radius: 10px;
            padding: 10px 16px;
            margin-bottom: 12px;
            box-shadow: var(--shadow);
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 12px;
            flex-wrap: wrap;
        }

        .brand-compact {
            font-size: 14px;
            font-weight: 800;
            color: #0f172a;
            letter-spacing: -0.3px;
            display: flex;
            align-items: center;
            gap: 8px;
        }

        .brand-dot {
            width: 8px;
            height: 8px;
            background: #10b981;
            border-radius: 50%;
        }

        .ctrl-group {
            display: flex;
            align-items: center;
            gap: 8px;
        }

        .ctrl-label {
            font-size: 11px;
            font-weight: 700;
            color: var(--text-sub);
            text-transform: uppercase;
            letter-spacing: 0.5px;
            white-space: nowrap;
        }

        .mode-pills {
            display: flex;
            background: #f1f5f9;
            padding: 3px;
            border-radius: 8px;
            gap: 3px;
        }

        .mode-pill {
            border: none;
            background: transparent;
            color: var(--text-sub);
            font-size: 12px;
            font-weight: 600;
            padding: 5px 12px;
            border-radius: 6px;
            cursor: pointer;
            transition: all 0.15s ease;
            white-space: nowrap;
            display: flex;
            align-items: center;
            gap: 5px;
        }

        .mode-pill.active {
            background: #ffffff;
            color: #047857;
            box-shadow: 0 1px 3px rgba(0,0,0,0.06);
            font-weight: 700;
        }

        .file-input-hidden { display: none !important; }

        .btn-file-select {
            background: #ffffff;
            border: 1px solid var(--border);
            color: var(--text-main);
            font-weight: 600;
            font-size: 12px;
            padding: 6px 14px;
            border-radius: 7px;
            cursor: pointer;
            transition: all 0.15s ease;
            display: flex;
            align-items: center;
            gap: 6px;
        }

        .btn-file-select:hover {
            background: #f8fafc;
            border-color: #cbd5e1;
        }

        .file-info-badge {
            font-size: 11.5px;
            font-weight: 500;
            color: var(--text-sub);
            background: #f1f5f9;
            padding: 5px 10px;
            border-radius: 6px;
            max-width: 220px;
            overflow: hidden;
            text-overflow: ellipsis;
            white-space: nowrap;
        }

        .btn-action-compact {
            background: #10b981;
            color: white;
            border: none;
            border-radius: 7px;
            padding: 6px 18px;
            font-size: 12.5px;
            font-weight: 700;
            cursor: pointer;
            display: flex;
            align-items: center;
            gap: 6px;
            transition: all 0.15s ease;
            box-shadow: 0 2px 6px rgba(16, 185, 129, 0.2);
            white-space: nowrap;
        }

        .btn-action-compact:hover:not(:disabled) {
            background: #059669;
            box-shadow: 0 4px 10px rgba(16, 185, 129, 0.3);
        }

        .timer-badge-compact {
            background: #ecfdf5;
            color: #047857;
            border: 1px solid #a7f3d0;
            padding: 4px 10px;
            border-radius: 6px;
            font-size: 12px;
            font-weight: 700;
            font-family: monospace;
            display: flex;
            align-items: center;
            gap: 5px;
        }

        /* 1VS1 RESIZABLE COMPARISON CONTAINER */
        .comparison-container {
            display: flex;
            align-items: stretch;
            gap: 0;
            height: calc(100vh - 180px);
            min-height: 560px;
            position: relative;
            width: 100%;
        }

        .resizer-gutter {
            width: 14px;
            margin: 0 4px;
            cursor: col-resize;
            display: flex;
            align-items: center;
            justify-content: center;
            user-select: none;
            z-index: 10;
            transition: background 0.2s ease;
            border-radius: 6px;
        }

        .resizer-gutter:hover, .resizer-gutter.dragging {
            background: rgba(37, 99, 235, 0.12);
        }

        .gutter-handle {
            width: 4px;
            height: 38px;
            background: #cbd5e1;
            border-radius: 2px;
            transition: all 0.2s ease;
        }

        .resizer-gutter:hover .gutter-handle, .resizer-gutter.dragging .gutter-handle {
            background: #2563eb;
            height: 52px;
            box-shadow: 0 0 8px rgba(37, 99, 235, 0.5);
        }

        .panel-card {
            background: var(--bg-card);
            border: 1px solid var(--border);
            border-radius: var(--radius);
            padding: 16px;
            box-shadow: var(--shadow);
            display: flex;
            flex-direction: column;
            height: 100%;
        }

        .panel-card-header {
            font-size: 14.5px;
            font-weight: 800;
            color: var(--text-main);
            margin-bottom: 12px;
            padding-bottom: 8px;
            border-bottom: 1px solid var(--border);
            display: flex;
            align-items: center;
            justify-content: space-between;
        }

        .original-viewer-box {
            background: #f8fafc;
            border: 1px solid var(--border);
            border-radius: 12px;
            padding: 16px;
            flex: 1;
            overflow-y: auto;
            scroll-behavior: smooth;
        }

        .section-label {
            font-size: 14px;
            font-weight: 700;
            color: var(--text-sub);
            text-transform: uppercase;
            letter-spacing: 0.5px;
            margin-bottom: 12px;
            display: block;
        }

        .mode-container {
            display: flex;
            flex-direction: column;
            gap: 12px;
            margin-bottom: 24px;
        }

        .mode-card {
            border: 2px solid var(--border);
            border-radius: 14px;
            padding: 16px;
            background: #ffffff;
            cursor: pointer;
            transition: all 0.2s ease;
            display: flex;
            align-items: flex-start;
            gap: 14px;
            user-select: none;
        }

        .mode-card:hover {
            border-color: #93c5fd;
            background: #f0f7ff;
        }

        .mode-card.active {
            border-color: var(--primary);
            background: var(--primary-light);
            box-shadow: 0 4px 14px rgba(37, 99, 235, 0.12);
        }

        .mode-card .icon {
            font-size: 26px;
            line-height: 1;
        }

        .mode-card .details {
            flex: 1;
        }

        .mode-card .title {
            font-size: 15px;
            font-weight: 700;
            color: var(--text-main);
            margin-bottom: 4px;
        }

        .mode-card.active .title {
            color: var(--primary);
        }

        .mode-card .desc {
            font-size: 12px;
            color: var(--text-sub);
            line-height: 1.45;
        }

        .drop-zone {
            border: 3px dashed #cbd5e1;
            border-radius: 14px;
            padding: 32px 20px;
            text-align: center;
            background: #f8fafc;
            cursor: pointer;
            transition: all 0.2s ease;
            margin-bottom: 24px;
            user-select: none;
        }

        .drop-zone:hover, .drop-zone.dragover {
            border-color: var(--primary);
            background: #eff6ff;
        }

        .drop-zone .icon {
            font-size: 44px;
            margin-bottom: 10px;
        }

        .drop-zone .prompt {
            font-weight: 700;
            font-size: 15px;
            color: var(--text-main);
            margin-bottom: 6px;
        }

        .drop-zone .formats {
            font-size: 12px;
            color: var(--text-sub);
        }

        .file-input-hidden {
            display: none !important;
        }

        .btn-action {
            width: 100%;
            padding: 16px 24px;
            font-size: 16px;
            font-weight: 700;
            border-radius: 12px;
            border: none;
            cursor: pointer;
            background: linear-gradient(135deg, #2563eb, #1d4ed8);
            color: #ffffff;
            box-shadow: 0 6px 18px rgba(37, 99, 235, 0.28);
            transition: all 0.2s ease;
            display: flex;
            align-items: center;
            justify-content: center;
            gap: 10px;
        }

        .btn-action:hover:not(:disabled) {
            background: linear-gradient(135deg, #1d4ed8, #1e40af);
            transform: translateY(-1px);
            box-shadow: 0 8px 22px rgba(37, 99, 235, 0.35);
        }

        .btn-action:disabled {
            background: #94a3b8;
            box-shadow: none;
            cursor: not-allowed;
        }

        /* OUTPUT TABS & PAGE NAV */
        .output-tabs {
            display: grid;
            grid-template-columns: repeat(3, 1fr);
            gap: 8px;
            background: #f1f5f9;
            padding: 6px;
            border-radius: 12px;
            margin-bottom: 14px;
        }

        .tab-btn {
            background: transparent;
            border: none;
            padding: 12px 10px;
            font-size: 13px;
            font-weight: 700;
            color: var(--text-sub);
            border-radius: 8px;
            cursor: pointer;
            transition: all 0.2s ease;
            text-align: center;
        }

        .tab-btn.active {
            background: #ffffff;
            color: var(--primary);
            box-shadow: 0 2px 6px rgba(0,0,0,0.06);
        }

        .page-nav-bar {
            display: none;
            align-items: center;
            gap: 8px;
            overflow-x: auto;
            padding: 8px 12px;
            background: #f8fafc;
            border: 1px solid var(--border);
            border-radius: 10px;
            margin-bottom: 14px;
            scrollbar-width: thin;
        }

        .page-btn {
            background: #ffffff;
            border: 1px solid var(--border);
            color: var(--text-main);
            font-size: 12px;
            font-weight: 600;
            padding: 6px 12px;
            border-radius: 6px;
            cursor: pointer;
            white-space: nowrap;
            transition: all 0.15s ease;
        }

        .page-btn:hover, .page-btn.active {
            background: var(--primary);
            color: #ffffff;
            border-color: var(--primary);
        }

        .tab-panel { display: none; }
        .tab-panel.active { display: block; }

        /* CUSTOM SCROLLBAR & SMOOTH CONTAINER */
        .preview-box {
            background: #ffffff;
            border: 1px solid var(--border);
            border-radius: 12px;
            padding: 26px;
            min-height: 540px;
            max-height: 720px;
            overflow-y: auto;
            scroll-behavior: smooth;
            line-height: 1.7;
            color: #1e293b;
        }

        .preview-box::-webkit-scrollbar, textarea.code-area::-webkit-scrollbar {
            width: 8px;
            height: 8px;
        }
        .preview-box::-webkit-scrollbar-track, textarea.code-area::-webkit-scrollbar-track {
            background: #f1f5f9;
            border-radius: 10px;
        }
        .preview-box::-webkit-scrollbar-thumb, textarea.code-area::-webkit-scrollbar-thumb {
            background: #cbd5e1;
            border-radius: 10px;
        }
        .preview-box::-webkit-scrollbar-thumb:hover, textarea.code-area::-webkit-scrollbar-thumb:hover {
            background: #94a3b8;
        }

        .preview-box h1 { font-size: 1.6rem; color: #1e3a8a; margin: 20px 0 10px; border-bottom: 2px solid #dbeafe; padding-bottom: 6px; }
        .preview-box h2 { font-size: 1.3rem; color: #1e40af; margin: 16px 0 8px; }
        .preview-box h3 { font-size: 1.1rem; color: #2563eb; margin: 14px 0 6px; }

        .preview-box table {
            width: 100%;
            border-collapse: separate;
            border-spacing: 0;
            margin: 20px 0;
            font-size: 0.94rem;
            border-radius: 10px;
            overflow: hidden;
            border: 1px solid #cbd5e1;
            box-shadow: 0 4px 12px rgba(0,0,0,0.03);
        }

        .preview-box th {
            background: linear-gradient(135deg, #1e40af, #2563eb);
            color: #ffffff;
            font-weight: 700;
            padding: 12px 16px;
            text-align: left;
            border-bottom: 2px solid #1d4ed8;
            letter-spacing: 0.3px;
        }

        .preview-box td {
            border-bottom: 1px solid #e2e8f0;
            border-right: 1px solid #f1f5f9;
            padding: 11px 16px;
            color: #1e293b;
            transition: background 0.15s ease;
        }

        .preview-box tr:last-child td {
            border-bottom: none;
        }

        .preview-box tr:nth-child(even) {
            background-color: #f8fafc;
        }

        .preview-box tr:hover td {
            background-color: #eff6ff;
        }

        .output-tabs-inline {
            display: flex;
            background: #f1f5f9;
            padding: 3px;
            border-radius: 8px;
            gap: 3px;
        }

        .tab-btn-sm {
            background: transparent;
            border: none;
            padding: 5px 12px;
            font-size: 12px;
            font-weight: 700;
            color: var(--text-sub);
            border-radius: 6px;
            cursor: pointer;
            transition: all 0.2s ease;
            white-space: nowrap;
        }

        .tab-btn-sm.active {
            background: #ffffff;
            color: var(--primary);
            box-shadow: 0 2px 5px rgba(0,0,0,0.06);
        }

        .export-actions-inline {
            display: flex;
            align-items: center;
            gap: 6px;
        }

        .btn-exp-sm {
            border: 1px solid var(--border);
            background: #ffffff;
            color: var(--text-main);
            font-size: 11.5px;
            font-weight: 700;
            padding: 5px 10px;
            border-radius: 7px;
            cursor: pointer;
            transition: all 0.2s ease;
            display: flex;
            align-items: center;
            gap: 4px;
            white-space: nowrap;
            box-shadow: 0 1px 2px rgba(0,0,0,0.04);
        }

        .btn-exp-sm:hover {
            transform: translateY(-1px);
            box-shadow: 0 3px 8px rgba(0,0,0,0.08);
        }

        .export-bar {
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 12px;
            background: #ffffff;
            border: 1px solid var(--border);
            padding: 10px 16px;
            border-radius: 12px;
            margin-bottom: 14px;
            box-shadow: 0 2px 6px rgba(0,0,0,0.02);
            flex-wrap: wrap;
        }

        .export-label {
            font-size: 13px;
            font-weight: 800;
            color: var(--text-main);
            display: flex;
            align-items: center;
            gap: 6px;
        }

        .export-actions {
            display: flex;
            align-items: center;
            gap: 8px;
            flex-wrap: wrap;
        }

        .btn-exp {
            border: 1px solid var(--border);
            background: #ffffff;
            color: var(--text-main);
            font-size: 12px;
            font-weight: 700;
            padding: 8px 14px;
            border-radius: 8px;
            cursor: pointer;
            transition: all 0.2s ease;
            display: flex;
            align-items: center;
            gap: 6px;
            box-shadow: 0 1px 3px rgba(0,0,0,0.04);
        }

        .btn-exp:hover {
            transform: translateY(-1px);
            box-shadow: 0 4px 10px rgba(0,0,0,0.08);
        }

        .btn-exp-docx {
            background: #eff6ff;
            color: #1d4ed8;
            border-color: #bfdbfe;
        }
        .btn-exp-docx:hover {
            background: #dbeafe;
            border-color: #93c5fd;
        }

        .btn-exp-pdf {
            background: #fef2f2;
            color: #dc2626;
            border-color: #fecaca;
        }
        .btn-exp-pdf:hover {
            background: #fee2e2;
            border-color: #fca5a5;
        }

        .btn-exp-txt {
            background: #f8fafc;
            color: #475569;
            border-color: #e2e8f0;
        }
        .btn-exp-txt:hover {
            background: #f1f5f9;
        }

        .btn-exp-copy {
            background: #f0fdf4;
            color: #166534;
            border-color: #bbf7d0;
        }
        .btn-exp-copy:hover {
            background: #dcfce7;
        }

        .page-break-divider {
            display: flex;
            align-items: center;
            gap: 12px;
            margin: 28px 0 16px;
            color: var(--primary);
            font-weight: 700;
            font-size: 13px;
        }

        .page-break-divider::before, .page-break-divider::after {
            content: '';
            flex: 1;
            height: 1px;
            background: #93c5fd;
        }

        textarea.code-area {
            width: 100%;
            height: 540px;
            background: #f8fafc;
            color: #0f172a;
            border: 1px solid var(--border);
            border-radius: 12px;
            padding: 18px;
            font-family: 'JetBrains Mono', Consolas, monospace;
            font-size: 13.5px;
            line-height: 1.6;
            resize: vertical;
            box-sizing: border-box;
            scroll-behavior: smooth;
        }

        .toolbar {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 10px;
        }

        .btn-copy {
            background: #ffffff;
            border: 1px solid var(--border);
            color: var(--text-main);
            font-weight: 600;
            padding: 8px 16px;
            border-radius: 8px;
            font-size: 13px;
            cursor: pointer;
            box-shadow: 0 1px 2px rgba(0,0,0,0.05);
            transition: all 0.2s ease;
        }

        .btn-copy:hover {
            background: #f1f5f9;
            border-color: #cbd5e1;
        }

        .btn-skip {
            display: none;
            background: #eff6ff;
            color: var(--primary);
            border: 1px solid #bfdbfe;
            font-weight: 600;
            padding: 6px 14px;
            border-radius: 8px;
            font-size: 12px;
            cursor: pointer;
        }

        .status-info {
            margin-top: 14px;
            text-align: center;
            font-size: 13px;
            font-weight: 600;
            color: var(--text-sub);
        }

        .spinner {
            display: none;
            width: 22px;
            height: 22px;
            border: 3px solid rgba(255,255,255,0.3);
            border-radius: 50%;
            border-top-color: white;
            animation: spin 0.8s linear infinite;
        }

        @keyframes spin { to { transform: rotate(360deg); } }

        @keyframes stitchBorderDraw {
            0% {
                box-shadow: 0 0 0 0 rgba(37, 99, 235, 0.4), inset 0 0 0 2px #3b82f6;
                border-color: #2563eb;
            }
            50% {
                box-shadow: 0 0 20px 4px rgba(6, 182, 212, 0.4), inset 0 0 12px 2px #06b6d4;
                border-color: #06b6d4;
            }
            100% {
                box-shadow: 0 4px 14px rgba(0, 0, 0, 0.05);
                border-color: #e2e8f0;
            }
        }

        .stitch-card {
            position: relative;
            border: 2px solid #3b82f6;
            border-radius: 12px;
            padding: 20px;
            margin-bottom: 20px;
            background: #ffffff;
            animation: stitchBorderDraw 2s ease-out forwards;
            overflow: hidden;
        }

        .stitch-laser-scan {
            position: absolute;
            top: 0;
            left: 0;
            width: 100%;
            height: 3px;
            background: linear-gradient(90deg, transparent, #06b6d4, #3b82f6, transparent);
            box-shadow: 0 0 10px #06b6d4;
            animation: laserScan 2.5s ease-in-out infinite;
            pointer-events: none;
        }

        @keyframes laserScan {
            0% { top: 0%; opacity: 1; }
            50% { top: 98%; opacity: 1; }
            100% { top: 0%; opacity: 0.2; }
        }
    </style>
</head>
<body>
    <main>
        <!-- Top Compact Control Toolbar (Notion Style) -->
        <div class="top-control-bar">
            <div class="brand-compact">
                <span class="brand-dot"></span>
                <span>chinhan_OCR</span>
            </div>

            <div class="ctrl-group">
                <div class="mode-pills">
                    <button class="mode-pill active" id="modeGundam" onclick="selectMode('gundam')">
                        <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><rect x="3" y="3" width="18" height="18" rx="2"/><line x1="9" y1="3" x2="9" y2="21"/><line x1="15" y1="3" x2="15" y2="21"/></svg>
                        Gundam (640px)
                    </button>
                    <button class="mode-pill" id="modeBase" onclick="selectMode('base')">
                        <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/></svg>
                        Base (1024px)
                    </button>
                </div>
            </div>

            <div class="ctrl-group">
                <input type="file" id="fileInput" class="file-input-hidden" accept="image/*,application/pdf,.pdf,.PDF" onchange="handleFileSelect(event)">
                <button class="btn-file-select" onclick="triggerFileInput()">
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="17 8 12 3 7 8"/><line x1="12" y1="3" x2="12" y2="15"/></svg>
                    Chọn tệp
                </button>
                <div class="file-info-badge" id="dropText">Chưa chọn tệp...</div>
            </div>

            <div class="ctrl-group">
                <button class="btn-action-compact" id="submitBtn" onclick="executeOCR()">
                    <div class="spinner" id="btnSpinner"></div>
                    <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><polygon points="5 3 19 12 5 21 5 3"/></svg>
                    <span id="btnText">Chạy</span>
                </button>
                <div class="timer-badge-compact" id="liveTimer">
                    <svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><polyline points="12 6 12 12 16 14"/></svg>
                    <span>00:00.0s</span>
                </div>
            </div>
            <div class="status-info" id="statusMsg" style="margin-top:0; font-size:12px; max-width:200px; overflow:hidden; text-overflow:ellipsis; white-space:nowrap;"></div>
        </div>

        <!-- Full Width Main Output Panel -->
        <div class="panel-card" style="height: calc(100vh - 120px); min-height: 600px;">
            <div class="panel-card-header" style="flex-wrap:wrap; gap:10px; margin-bottom:12px; padding-bottom:10px;">
                <div style="display:flex; align-items:center; gap:14px;">
                    <span style="font-size:15px; font-weight:700; color:#0f172a;">Kết quả trích xuất</span>
                    <!-- 3tab Navigation Inline -->
                    <div class="output-tabs-inline">
                        <button class="tab-btn-sm active" id="tabBtnPreview" onclick="switchOutputTab('preview')">
                            <svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M12 2l3.09 6.26L22 9.27l-5 4.87 1.18 6.88L12 17.77l-6.18 3.25L7 14.14 2 9.27l6.91-1.01L12 2z"/></svg>
                            Blueprint
                        </button>
                        <button class="tab-btn-sm" id="tabBtnClean" onclick="switchOutputTab('clean')">
                            <svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="16" y1="13" x2="8" y2="13"/><line x1="16" y1="17" x2="8" y2="17"/></svg>
                            Markdown
                        </button>
                        <button class="tab-btn-sm" id="tabBtnRaw" onclick="switchOutputTab('raw')">
                            <svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><polyline points="16 18 22 12 16 6"/><polyline points="8 6 2 12 8 18"/></svg>
                            Raw
                        </button>
                    </div>
                </div>

                <!-- Export & Import Actions Inline -->
                <div class="export-actions-inline">
                    <input type="file" id="importFileInput" class="file-input-hidden" accept=".docx,.pdf,.txt,.md" onchange="handleImportFileSelect(event)">
                    <button class="btn-exp-sm" style="background:#ffffff; color:#374151; border-color:#d1d5db;" onclick="triggerImportFileInput()" title="Import tệp Word/PDF/TXT/MD">
                        <svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>
                        Import
                    </button>
                    <span style="font-size:11px; font-weight:700; color:var(--text-sub); margin-left:4px;">Tải về:</span>
                    <button class="btn-exp-sm btn-exp-docx" onclick="downloadDOCX()">Word (.docx)</button>
                    <button class="btn-exp-sm btn-exp-pdf" onclick="downloadPDF()">PDF (.pdf)</button>
                    <button class="btn-exp-sm btn-exp-md" onclick="downloadMD()" style="background:#ffffff; color:#047857; border-color:#a7f3d0;">Markdown (.md)</button>
                    <button class="btn-exp-sm btn-exp-txt" onclick="downloadTXT()">TXT (.txt)</button>
                    <button class="btn-exp-sm btn-exp-copy" onclick="copyCurrentTabContent()">Copy</button>
                </div>
            </div>

            <!-- Sticky Page Nav Bar for PDF Multi-page -->
            <div class="page-nav-bar" id="pageNavBar"></div>

            <!-- Tab 1: Render View -->
            <div class="tab-panel active" id="panelPreview" style="flex:1; display:flex; flex-direction:column; overflow:hidden;">
                <div class="preview-box" id="previewArea" style="flex:1; max-height:none;">Tài liệu đã được trích xuất và định dạng trình bày Y Chang Ảnh sẽ hiển thị ở đây...</div>
            </div>

            <!-- Tab 2: Clean Markdown -->
            <div class="tab-panel" id="panelClean" style="flex:1; display:none; flex-direction:column; overflow:hidden;">
                <textarea class="code-area" id="cleanText" readonly style="flex:1; height:100%;" placeholder="Kết quả Markdown sạch (dùng để dán vào Word/Notion) sẽ hiển thị ở đây..."></textarea>
            </div>

            <!-- Tab 3: Raw Bounding Box -->
            <div class="tab-panel" id="panelRaw" style="flex:1; display:none; flex-direction:column; overflow:hidden;">
                <textarea class="code-area" id="rawText" readonly style="flex:1; height:100%;" placeholder="Kết quả OCR gốc chứa các thẻ tọa độ <|det|>..."></textarea>
            </div>
        </div>
    </main>

    <script>
        let currentMode = 'gundam';
        let currentFile = null;
        let timerInterval = null;
        let startTime = 0;

        function selectMode(mode) {
            currentMode = mode;
            document.getElementById('modeGundam').classList.toggle('active', mode === 'gundam');
            document.getElementById('modeBase').classList.toggle('active', mode === 'base');
        }

        function triggerFileInput() {
            document.getElementById('fileInput').click();
        }

        function handleFileSelect(e) {
            if (e.target && e.target.files && e.target.files.length > 0) {
                setFile(e.target.files[0]);
            }
        }

        const dropZone = document.getElementById('dropZone');
        if (dropZone) {
            dropZone.addEventListener('dragover', (e) => { e.preventDefault(); e.stopPropagation(); dropZone.classList.add('dragover'); });
            dropZone.addEventListener('dragleave', (e) => { e.preventDefault(); e.stopPropagation(); dropZone.classList.remove('dragover'); });
            dropZone.addEventListener('drop', (e) => {
                e.preventDefault();
                e.stopPropagation();
                dropZone.classList.remove('dragover');
                if (e.dataTransfer && e.dataTransfer.files && e.dataTransfer.files.length > 0) {
                    setFile(e.dataTransfer.files[0]);
                }
            });
        }

        if (typeof pdfjsLib !== 'undefined') {
            pdfjsLib.GlobalWorkerOptions.workerSrc = 'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js';
        }

        function setFile(file) {
            currentFile = file;
            const sizeMB = (file.size / (1024 * 1024)).toFixed(2);
            const icon = file.type.includes('pdf') || file.name.endsWith('.pdf') ? '📄' : '🖼️';
            document.getElementById('dropText').innerText = `${icon} ${file.name} (${sizeMB} MB)`;
        }

        function switchOutputTab(tab) {
            const p1 = document.getElementById('panelPreview');
            const p2 = document.getElementById('panelClean');
            const p3 = document.getElementById('panelRaw');

            document.getElementById('tabBtnPreview').classList.toggle('active', tab === 'preview');
            document.getElementById('tabBtnClean').classList.toggle('active', tab === 'clean');
            document.getElementById('tabBtnRaw').classList.toggle('active', tab === 'raw');

            if (p1) {
                p1.classList.toggle('active', tab === 'preview');
                p1.style.display = (tab === 'preview') ? 'flex' : 'none';
            }
            if (p2) {
                p2.classList.toggle('active', tab === 'clean');
                p2.style.display = (tab === 'clean') ? 'flex' : 'none';
            }
            if (p3) {
                p3.classList.toggle('active', tab === 'raw');
                p3.style.display = (tab === 'raw') ? 'flex' : 'none';
            }
        }

        window.currentCleanMarkdown = "";

        function triggerImportFileInput() {
            const input = document.getElementById('importFileInput');
            if (input) input.click();
        }

        async function handleImportFileSelect(e) {
            if (!e.target || !e.target.files || e.target.files.length === 0) return;
            const file = e.target.files[0];
            const statusMsg = document.getElementById('statusMsg');
            statusMsg.innerText = `⏳ Đang import tệp ${file.name}...`;

            const formData = new FormData();
            formData.append('file', file);

            try {
                const response = await fetch('/v1/import/file', {
                    method: 'POST',
                    body: formData
                });
                if (!response.ok) {
                    const err = await response.json();
                    alert("Lỗi import: " + (err.detail || "Không thể nạp tệp"));
                    statusMsg.innerText = "❌ Lỗi import tệp.";
                    return;
                }
                const data = await response.json();
                window.currentCleanMarkdown = data.markdown_text;

                document.getElementById('cleanText').value = data.markdown_text;
                document.getElementById('rawText').value = data.raw_text;
                document.getElementById('previewArea').innerHTML = `
                    <div class="stitch-card" style="animation:none;">
                        <div class="page-break-divider">📄 TỆP IMPORT: ${data.filename}</div>
                        <div class="stitch-content">${marked.parse(data.markdown_text)}</div>
                    </div>
                `;
                statusMsg.innerText = `✨ Đã import thành công ${data.filename} vào Màn 2!`;
            } catch (err) {
                alert("Lỗi kết nối import: " + err.message);
                statusMsg.innerText = "❌ Lỗi import tệp.";
            } finally {
                e.target.value = '';
            }
        }

        async function downloadDOCX() {
            const markdown = window.currentCleanMarkdown || document.getElementById('cleanText').value || document.getElementById('previewArea').innerText;
            if (!markdown || !markdown.trim() || markdown.includes("Tài liệu đã được trích xuất")) {
                alert("Chưa có dữ liệu trích xuất để tạo file Word (.docx)!");
                return;
            }
            const baseName = currentFile ? currentFile.name.replace(/\.[^/.]+$/, "") : "ocr_result";
            try {
                const response = await fetch('/v1/export/docx', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ markdown_text: markdown, filename: baseName + ".docx" })
                });
                if (!response.ok) {
                    alert("Không thể tạo tệp Word trên server.");
                    return;
                }
                const blob = await response.blob();
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = baseName + ".docx";
                document.body.appendChild(a);
                a.click();
                a.remove();
                window.URL.revokeObjectURL(url);
            } catch (err) {
                alert("Lỗi tải file Word (.docx): " + err.message);
            }
        }

        function downloadPDF() {
            const previewArea = document.getElementById('previewArea');
            if (!previewArea || !previewArea.innerText.trim() || previewArea.innerText.includes("Tài liệu đã được trích xuất")) {
                alert("Chưa có dữ liệu trích xuất để tạo file PDF!");
                return;
            }
            const baseName = currentFile ? currentFile.name.replace(/\.[^/.]+$/, "") : "ocr_result";

            const clone = previewArea.cloneNode(true);
            clone.style.position = 'absolute';
            clone.style.left = '-9999px';
            clone.style.top = '0';
            clone.style.width = '800px';
            clone.style.height = 'auto';
            clone.style.overflow = 'visible';
            clone.style.background = '#ffffff';
            clone.style.padding = '20px';
            document.body.appendChild(clone);

            const opt = {
                margin:       [0.4, 0.4, 0.4, 0.4],
                filename:     baseName + ".pdf",
                image:        { type: 'jpeg', quality: 0.98 },
                html2canvas:  { scale: 2, useCORS: true, scrollY: 0 },
                jsPDF:        { unit: 'in', format: 'a4', orientation: 'portrait' }
            };

            if (typeof html2pdf !== 'undefined') {
                html2pdf().set(opt).from(clone).save().then(() => {
                    if (clone.parentNode) document.body.removeChild(clone);
                }).catch(err => {
                    if (clone.parentNode) document.body.removeChild(clone);
                    alert("Lỗi tạo file PDF: " + err.message);
                });
            } else {
                if (clone.parentNode) document.body.removeChild(clone);
                window.print();
            }
        }

        function downloadTXT() {
            const content = window.currentCleanMarkdown || document.getElementById('cleanText').value || document.getElementById('rawText').value || document.getElementById('previewArea').innerText;
            if (!content || !content.trim() || content.includes("Tài liệu đã được trích xuất")) {
                alert("Chưa có dữ liệu để tải về!");
                return;
            }
            const baseName = currentFile ? currentFile.name.replace(/\.[^/.]+$/, "") : "ocr_result";
            const blob = new Blob([content], { type: 'text/plain;charset=utf-8' });
            const url = window.URL.createObjectURL(blob);
            const a = document.createElement('a');
            a.href = url;
            a.download = baseName + ".txt";
            document.body.appendChild(a);
            a.click();
            a.remove();
            window.URL.revokeObjectURL(url);
        }

        function downloadMD() {
            const content = window.currentCleanMarkdown || document.getElementById('cleanText').value || document.getElementById('previewArea').innerText;
            if (!content || !content.trim() || content.includes("Tài liệu đã được trích xuất")) {
                alert("Chưa có dữ liệu Markdown để tải về!");
                return;
            }
            const baseName = currentFile ? currentFile.name.replace(/\.[^/.]+$/, "") : "ocr_result";
            const blob = new Blob([content], { type: 'text/markdown;charset=utf-8' });
            const url = window.URL.createObjectURL(blob);
            const a = document.createElement('a');
            a.href = url;
            a.download = baseName + ".md";
            document.body.appendChild(a);
            a.click();
            a.remove();
            window.URL.revokeObjectURL(url);
        }

        function copyCurrentTabContent() {
            let textToCopy = '';
            const previewPanel = document.getElementById('panelPreview');
            const cleanPanel = document.getElementById('panelClean');
            const rawPanel = document.getElementById('panelRaw');

            if (previewPanel.classList.contains('active')) {
                textToCopy = window.currentCleanMarkdown || document.getElementById('previewArea').innerText;
            } else if (cleanPanel.classList.contains('active')) {
                textToCopy = document.getElementById('cleanText').value;
            } else if (rawPanel.classList.contains('active')) {
                textToCopy = document.getElementById('rawText').value;
            }

            if (!textToCopy || !textToCopy.trim() || textToCopy.includes("Tài liệu đã được trích xuất")) {
                alert("Không có nội dung để sao chép!");
                return;
            }

            navigator.clipboard.writeText(textToCopy).then(() => {
                alert("Đã sao chép thành công!");
            }).catch(err => {
                alert("Lỗi sao chép: " + err);
            });
        }

        function startTimer() {
            startTime = Date.now();
            clearInterval(timerInterval);
            const timerEl = document.getElementById('liveTimer');
            timerInterval = setInterval(() => {
                const elapsedSec = ((Date.now() - startTime) / 1000).toFixed(1);
                if (timerEl) {
                    timerEl.innerHTML = `<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><polyline points="12 6 12 12 16 14"/></svg><span>${elapsedSec}s</span>`;
                }
            }, 100);
        }

        function stopTimer() {
            clearInterval(timerInterval);
            const totalTime = ((Date.now() - startTime) / 1000).toFixed(2);
            const timerEl = document.getElementById('liveTimer');
            if (timerEl) {
                timerEl.innerHTML = `<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><polyline points="12 6 12 12 16 14"/></svg><span>${totalTime}s</span>`;
            }
            return totalTime;
        }

        async function executeOCR() {
            if (!currentFile) {
                alert("Vui lòng chọn tệp Ảnh hoặc tệp PDF trước khi bấm trích xuất!");
                return;
            }

            const btn = document.getElementById('submitBtn');
            const btnText = document.getElementById('btnText');
            const btnSpinner = document.getElementById('btnSpinner');
            const statusMsg = document.getElementById('statusMsg');

            btn.disabled = true;
            btnSpinner.style.display = 'inline-block';
            btnText.innerText = 'Đang chạy...';
            statusMsg.innerText = 'Đang trích xuất dữ liệu...';

            startTimer();

            const formData = new FormData();
            formData.append('file', currentFile);

            const previewArea = document.getElementById('previewArea');
            const cleanTextArea = document.getElementById('cleanText');
            const rawTextArea = document.getElementById('rawText');

            window.currentCleanMarkdown = '';
            previewArea.innerHTML = '';
            cleanTextArea.value = '';
            rawTextArea.value = '';

            try {
                const response = await fetch(`/v1/ocr/stream?mode=${currentMode}`, {
                    method: 'POST',
                    body: formData
                });

                if (!response.ok) {
                    const errText = await response.text();
                    alert("Lỗi Server: " + errText);
                    statusMsg.innerText = "Có lỗi xảy ra khi thực thi OCR Stream.";
                    return;
                }

                const reader = response.body.getReader();
                const decoder = new TextDecoder('utf-8');
                let buffer = '';

                while (true) {
                    const { done, value } = await reader.read();
                    if (done) break;
                    buffer += decoder.decode(value, { stream: true });

                    const eventBlocks = buffer.split('\n\n');
                    buffer = eventBlocks.pop();

                    for (const block of eventBlocks) {
                        if (!block.trim()) continue;
                        const lines = block.split('\n');
                        let eventName = '';
                        let dataStr = '';

                        for (const line of lines) {
                            if (line.startsWith('event: ')) eventName = line.slice(7).trim();
                            if (line.startsWith('data: ')) dataStr = line.slice(6).trim();
                        }

                        if (eventName === 'page_data' && dataStr) {
                            const pageObj = JSON.parse(dataStr);
                            appendPageStitchUI(pageObj);
                        } else if (eventName === 'complete') {
                            const totalSec = stopTimer();
                            statusMsg.innerText = `Hoàn tất trong ${totalSec}s`;
                        }
                    }
                }
            } catch (err) {
                stopTimer();
                alert("Lỗi kết nối Server Stream: " + err.message);
            } finally {
                btn.disabled = false;
                btnSpinner.style.display = 'none';
                btnText.innerText = 'Chạy';
            }
        }

        function appendPageStitchUI(pageObj) {
            const previewArea = document.getElementById('previewArea');
            const cleanTextArea = document.getElementById('cleanText');
            const rawTextArea = document.getElementById('rawText');

            window.currentCleanMarkdown += (window.currentCleanMarkdown ? '\n\n' : '') + pageObj.clean_markdown;

            // Update Page Nav Buttons Real-time
            updatePageNavBar(pageObj.page_index, pageObj.total_pages);

            // Tab 1: Stitch Blueprint Card
            const card = document.createElement('div');
            card.className = 'stitch-card';
            card.id = `page-card-${pageObj.page_index}`;
            card.setAttribute('data-page', pageObj.page_index);
            card.innerHTML = `
                <div class="stitch-laser-scan"></div>
                <div class="page-break-divider">📄 TRANG ${pageObj.page_index} / ${pageObj.total_pages} (${pageObj.elapsed_seconds}s)</div>
                <div class="stitch-content" id="stitch-content-${pageObj.page_index}"></div>
            `;
            previewArea.appendChild(card);
            previewArea.scrollTop = previewArea.scrollHeight;

            // Line-by-Line Typewriter Stream Rendering matching GitHub demo GIF
            const contentEl = card.querySelector(`#stitch-content-${pageObj.page_index}`);
            const lines = pageObj.clean_markdown.split('\n');
            let lineIdx = 0;
            let accumulatedMd = '';

            const streamTimer = setInterval(() => {
                if (lineIdx < lines.length) {
                    const step = lines.length > 30 ? 3 : (lines.length > 15 ? 2 : 1);
                    for (let s = 0; s < step && lineIdx < lines.length; s++) {
                        accumulatedMd += lines[lineIdx] + '\n';
                        lineIdx++;
                    }
                    contentEl.innerHTML = marked.parse(accumulatedMd);
                    previewArea.scrollTop = previewArea.scrollHeight;
                } else {
                    clearInterval(streamTimer);
                    setTimeout(() => {
                        const laser = card.querySelector('.stitch-laser-scan');
                        if (laser) laser.remove();
                    }, 800);
                }
            }, 10);

            // Tab 2 & 3: Append text
            cleanTextArea.value += `\n*— Trang ${pageObj.page_index} —*\n` + pageObj.clean_markdown + '\n';
            cleanTextArea.scrollTop = cleanTextArea.scrollHeight;

            rawTextArea.value += `--- Page ${pageObj.page_index} ---\n` + pageObj.raw_text + '\n';
            rawTextArea.scrollTop = rawTextArea.scrollHeight;
        }

        function updatePageNavBar(pageIndex, totalPages) {
            const pageNavBar = document.getElementById('pageNavBar');
            if (!pageNavBar) return;

            if (totalPages > 1) {
                pageNavBar.style.display = 'flex';
            }

            if (!document.getElementById('page-nav-btn-all')) {
                pageNavBar.innerHTML = '';
                const allBtn = document.createElement('button');
                allBtn.id = 'page-nav-btn-all';
                allBtn.className = 'page-btn active';
                allBtn.innerText = `📄 Xem Tất Cả (${totalPages} trang)`;
                allBtn.onclick = () => filterPageCard(0);
                pageNavBar.appendChild(allBtn);
            }

            if (!document.getElementById(`page-nav-btn-${pageIndex}`)) {
                const btn = document.createElement('button');
                btn.id = `page-nav-btn-${pageIndex}`;
                btn.className = 'page-btn';
                btn.innerText = `Trang ${pageIndex}`;
                btn.onclick = () => filterPageCard(pageIndex);
                pageNavBar.appendChild(btn);
            }
        }

        function filterPageCard(targetPageIndex) {
            document.querySelectorAll('.page-btn').forEach(b => b.classList.remove('active'));
            if (targetPageIndex === 0) {
                const allBtn = document.getElementById('page-nav-btn-all');
                if (allBtn) allBtn.classList.add('active');
                document.querySelectorAll('.stitch-card').forEach(card => {
                    card.style.display = 'block';
                });
                document.getElementById('previewArea').scrollTop = 0;
            } else {
                const btn = document.getElementById(`page-nav-btn-${targetPageIndex}`);
                if (btn) btn.classList.add('active');
                document.querySelectorAll('.stitch-card').forEach(card => {
                    if (card.id === `page-card-${targetPageIndex}`) {
                        card.style.display = 'block';
                        card.scrollIntoView({ behavior: 'smooth', block: 'start' });
                    } else {
                        card.style.display = 'none';
                    }
                });
            }
        }

        function formatMarkdownWithPageDividers(text) {
            let pageIndex = 1;
            let formatted = text.replace(/(\*— Trang \d+ —\*|--- Page \d+ ---)/g, (match) => {
                const id = `page-heading-${pageIndex}`;
                const badge = `<div class="page-break-divider" id="${id}">📄 TRANG ${pageIndex}</div>`;
                pageIndex++;
                return badge;
            });
            return formatted;
        }

        function copyCleanText() {
            const cleanText = document.getElementById('cleanText');
            cleanText.select();
            navigator.clipboard.writeText(cleanText.value);
            alert("Đã sao chép Markdown Sạch vào bộ nhớ tạm!");
        }

        function copyRawText() {
            const rawText = document.getElementById('rawText');
            rawText.select();
            navigator.clipboard.writeText(rawText.value);
            alert("Đã sao chép Dữ Liệu Gốc vào bộ nhớ tạm!");
        }
    </script>
</body>
</html>"""
        return HTMLResponse(content=html_content)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=False)
